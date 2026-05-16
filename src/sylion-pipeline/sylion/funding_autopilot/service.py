from __future__ import annotations

import csv
import json
import time
import zipfile
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape

from .config import funding_results_root
from .governance_bridge import check_approved, submit_submission_ticket
from .official_registry import fetch_krs_company_profile
from .store import get_funding_store


CORE_PROFILE_FIELDS = [
    "legal_name",
    "tax_id",
    "country",
    "legal_form",
    "sme_status",
    "employees",
    "annual_revenue",
    "technologies",
    "products",
    "representative_name",
    "representative_email",
]

CORE_DOCUMENT_TYPES = [
    "financial_statement",
    "tax_clearance",
    "social_security_clearance",
    "incorporation_document",
]

EXPORT_ARTIFACT_MEDIA_TYPES = {
    "json": "application/json",
    "markdown": "text/markdown; charset=utf-8",
    "csv": "text/csv; charset=utf-8",
    "review": "application/json",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "pdf": "application/pdf",
    "zip": "application/zip",
}

EXPORT_SECTION_LABELS = {
    "executive_summary": "Podsumowanie wykonawcze",
    "project_description": "Opis projektu",
    "innovation_case": "Uzasadnienie innowacji",
    "market_case": "Uzasadnienie rynkowe",
    "implementation_plan": "Plan wdrozenia",
    "budget": "Budzet",
    "consortium": "Konsorcjum",
    "compliance": "Zgodnosc formalna",
}

EXPORT_BUDGET_LABELS = {
    "budget_total": "Budzet calkowity",
    "grant_requested": "Wnioskowany grant",
    "own_contribution": "Wklad wlasny",
    "grant_intensity_pct": "Intensywnosc dofinansowania (%)",
    "cost_categories": "Kategorie kosztow",
}

CATEGORY_RULES = [
    ("projekt AI", {"ai", "ml", "machine learning", "data", "vision"}),
    ("projekt cyberbezpieczenstwa", {"security", "cyber", "soc", "siem", "zero trust"}),
    ("projekt energetyczny", {"energy", "oze", "solar", "battery", "grid"}),
    ("projekt automatyzacji", {"automation", "robot", "robotics", "industry 4.0"}),
    ("projekt eksportowy", {"export", "international", "market expansion"}),
]


def _now() -> float:
    return time.time()


def _export_section_label(section_name: str) -> str:
    return EXPORT_SECTION_LABELS.get(section_name, section_name.replace("_", " ").title())


def _export_budget_label(metric: str) -> str:
    return EXPORT_BUDGET_LABELS.get(metric, metric.replace("_", " ").title())


def _xlsx_inline_cell(column: str, row_number: int, value: Any) -> str:
    text = xml_escape("" if value is None else str(value))
    return f'<c r="{column}{row_number}" t="inlineStr"><is><t>{text}</t></is></c>'


def _write_budget_xlsx(path: Path, budget: dict[str, Any]) -> None:
    rows = [("Metryka", "Wartosc")]
    rows.extend((_export_budget_label(key), value) for key, value in budget.items())
    sheet_rows = []
    for row_number, (metric, value) in enumerate(rows, start=1):
        sheet_rows.append(
            f'<row r="{row_number}">'
            f'{_xlsx_inline_cell("A", row_number, metric)}'
            f'{_xlsx_inline_cell("B", row_number, value)}'
            "</row>"
        )
    worksheet = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        "<sheetData>"
        + "".join(sheet_rows)
        + "</sheetData></worksheet>"
    )
    workbook = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<sheets><sheet name="Budzet" sheetId="1" r:id="rId1"/></sheets></workbook>'
    )
    workbook_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
        'Target="worksheets/sheet1.xml"/></Relationships>'
    )
    root_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="xl/workbook.xml"/>'
        '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" '
        'Target="docProps/core.xml"/>'
        '<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" '
        'Target="docProps/app.xml"/></Relationships>'
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
        '<Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
        '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
        '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>'
        "</Types>"
    )
    core = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
        'xmlns:dc="http://purl.org/dc/elements/1.1/" '
        'xmlns:dcterms="http://purl.org/dc/terms/" '
        'xmlns:dcmitype="http://purl.org/dc/dcmitype/" '
        'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
        "<dc:title>Budzet wniosku grantowego</dc:title>"
        "<dc:creator>AEIS Funding Autopilot</dc:creator>"
        "</cp:coreProperties>"
    )
    app = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" '
        'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
        "<Application>AEIS Funding Autopilot</Application></Properties>"
    )
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", root_rels)
        archive.writestr("xl/workbook.xml", workbook)
        archive.writestr("xl/_rels/workbook.xml.rels", workbook_rels)
        archive.writestr("xl/worksheets/sheet1.xml", worksheet)
        archive.writestr("docProps/core.xml", core)
        archive.writestr("docProps/app.xml", app)


def _pdf_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _write_text_pdf(path: Path, lines: list[str]) -> None:
    safe_lines = [line.encode("latin-1", "replace").decode("latin-1") for line in lines]
    stream_lines = ["BT", "/F1 11 Tf", "40 800 Td", "14 TL"]
    for line in safe_lines:
        stream_lines.append(f"({_pdf_escape(line[:120])}) Tj")
        stream_lines.append("T*")
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = []
    for index, payload in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode("ascii"))
        output.extend(payload)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii"))
    for offset in offsets:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    path.write_bytes(bytes(output))


def _clean_words(values: list[str]) -> set[str]:
    words: set[str] = set()
    for value in values:
        for token in str(value or "").replace("/", " ").replace(",", " ").lower().split():
            token = token.strip()
            if len(token) >= 2:
                words.add(token)
    return words


def _expand_query_tokens(tokens: set[str]) -> set[str]:
    expanded = set(tokens)
    synonym_groups = [
        {
            "ledger",
            "reconciliation",
            "fraud",
            "banking",
            "open-banking",
            "openbanking",
            "finance",
            "financial",
            "fintech",
            "cybersecurity",
            "cyberbezpieczenstwo",
            "cyberbezpieczeństwo",
            "kryptografia",
            "cryptography",
            "compliance",
            "ai",
        },
        {"qkd", "quantum", "post-quantum", "postkwantowa", "kryptografia", "cryptography"},
        {"energy", "grid", "energia", "oze", "smart", "efektywność", "efficiency"},
        {"esg", "csrd", "carbon", "emission", "supplier", "supply-chain", "greenwashing"},
    ]
    for group in synonym_groups:
        if expanded & group:
            expanded.update(group)
    return expanded


def _keywords_for_company(profile: dict[str, Any]) -> set[str]:
    return _clean_words(
        list(profile.get("technologies", []))
        + list(profile.get("products", []))
        + list(profile.get("services", []))
        + list(profile.get("strategic_goals", []))
        + list(profile.get("team_competencies", []))
        + list(profile.get("export_markets", []))
    )


def _keywords_for_call(call: dict[str, Any]) -> set[str]:
    return _clean_words(
        list(call.get("themes_json", []))
        + list(call.get("required_partner_types_json", []))
        + [call.get("title", ""), call.get("code", ""), call.get("country", ""), call.get("region", "")]
    )


def _keywords_for_project(project: dict[str, Any]) -> set[str]:
    return _clean_words(
        [project.get("title", ""), project.get("summary", ""), project.get("objective", "")]
        + list(project.get("target_markets", []))
        + list(project.get("partner_needs", []))
        + list(project.get("work_packages", []))
    )


def _overlap_score(lhs: set[str], rhs: set[str]) -> float:
    if not lhs or not rhs:
        return 0.0
    intersection = len(lhs & rhs)
    union = len(lhs | rhs)
    return round((intersection / max(union, 1)) * 100, 2)


def _clamp(value: float, low: float = 0.0, high: float = 100.0) -> float:
    return round(max(low, min(high, value)), 2)


def _country_key(value: Any) -> str:
    normalized = str(value or "").strip().lower()
    normalized = normalized.replace("ą", "a").replace("ł", "l").replace("ń", "n").replace("ó", "o").replace("ś", "s").replace("ż", "z").replace("ź", "z").replace("ę", "e").replace("ć", "c")
    aliases = {
        "polska": "pl",
        "poland": "pl",
        "pl": "pl",
        "eu": "eu",
        "ue": "eu",
        "european union": "eu",
        "unia europejska": "eu",
    }
    return aliases.get(normalized, normalized)


def _country_compatible(company_country: Any, call_country: Any) -> bool:
    company = _country_key(company_country)
    call = _country_key(call_country)
    return call in {"", "eu"} or company == call


def _beneficiary_key(value: Any) -> str:
    normalized = str(value or "").strip().lower()
    normalized = normalized.replace("mśp", "sme").replace("msp", "sme")
    normalized = normalized.replace("małe i średnie przedsiębiorstwa", "sme")
    return normalized


def _dedupe_preserve_order(values: list[str]) -> list[str]:
    seen: set[str] = set()
    unique: list[str] = []
    for value in values:
        item = str(value or "").strip()
        if not item or item in seen:
            continue
        seen.add(item)
        unique.append(item)
    return unique


def _categorize(profile: dict[str, Any], call: dict[str, Any] | None = None) -> str:
    keywords = _keywords_for_company(profile)
    if call:
        keywords |= _keywords_for_call(call)
    for label, triggers in CATEGORY_RULES:
        if keywords & triggers:
            return label
    return "projekt B+R"


def _company_readiness(profile: dict[str, Any], documents: list[dict[str, Any]]) -> dict[str, Any]:
    missing_fields = []
    completed_fields = 0
    for field in CORE_PROFILE_FIELDS:
        value = profile.get(field)
        if isinstance(value, list):
            if value:
                completed_fields += 1
            else:
                missing_fields.append(field)
        elif value not in (None, "", 0):
            completed_fields += 1
        else:
            missing_fields.append(field)

    available_docs = {item.get("document_type", "") for item in documents if item.get("status") == "available"}
    missing_documents = [doc for doc in CORE_DOCUMENT_TYPES if doc not in available_docs]
    doc_score = 100.0 * ((len(CORE_DOCUMENT_TYPES) - len(missing_documents)) / max(len(CORE_DOCUMENT_TYPES), 1))
    field_score = 100.0 * (completed_fields / max(len(CORE_PROFILE_FIELDS), 1))
    readiness_score = _clamp((field_score * 0.7) + (doc_score * 0.3))
    return {
        "readiness_score": readiness_score,
        "missing_fields": missing_fields,
        "missing_documents": missing_documents,
        "field_score": round(field_score, 2),
        "document_score": round(doc_score, 2),
        "recommended_next_steps": [
            *[f"Uzupelnij pole: {field}" for field in missing_fields[:5]],
            *[f"Dostarcz dokument: {doc}" for doc in missing_documents[:5]],
        ],
    }


class FundingAutopilotService:
    def __init__(self, db_path: str | None = None) -> None:
        self.store = get_funding_store(db_path)

    def _empty_company_profile(self, company_id: str) -> dict[str, Any]:
        return {
            "company_id": company_id,
            "legal_name": "",
            "tax_id": "",
            "registration_id": "",
            "eu_vat": "",
            "country": "Poland",
            "region": "",
            "city": "",
            "legal_form": "",
            "established_at": "",
            "sme_status": "SME",
            "employees": 0,
            "annual_revenue": 0.0,
            "ebitda": 0.0,
            "total_assets": 0.0,
            "total_liabilities": 0.0,
            "export_markets": [],
            "technologies": [],
            "products": [],
            "services": [],
            "certifications": [],
            "patents": [],
            "team_competencies": [],
            "strategic_goals": [],
            "representative_name": "",
            "representative_email": "",
            "state_aid_total_eur": 0.0,
            "de_minimis_total_eur": 0.0,
            "prior_grants_count": 0,
            "notes": "",
        }

    def _require_company(self, company_id: str) -> dict[str, Any]:
        profile = self.store.get_company_profile(company_id)
        if not profile:
            raise ValueError(f"Company profile '{company_id}' not found")
        return profile

    def _require_project(self, project_id: str) -> dict[str, Any]:
        project = self.store.get_project(project_id)
        if not project:
            raise ValueError(f"Funding project '{project_id}' not found")
        return project

    def _require_call(self, call_id: str) -> dict[str, Any]:
        call = self.store.get_call(call_id)
        if not call:
            raise ValueError(f"Funding call '{call_id}' not found")
        return call

    def _require_application(self, application_id: str) -> dict[str, Any]:
        application = self.store.get_application(application_id)
        if not application:
            raise ValueError(f"Funding application '{application_id}' not found")
        return application

    def _require_submission_session(self, session_id: str) -> dict[str, Any]:
        session = self.store.get_submission_session(session_id)
        if not session:
            raise ValueError(f"Funding submission session '{session_id}' not found")
        return session

    def _application_document_state(self, company_id: str, required_documents: list[str]) -> tuple[list[str], list[str]]:
        company_documents = self.store.list_company_documents(company_id)
        available_documents = _dedupe_preserve_order(
            [item.get("document_type", "") for item in company_documents if item.get("status") == "available"]
        )
        missing_documents = [item for item in required_documents if item not in available_documents]
        return available_documents, missing_documents

    def _status_after_document_refresh(self, application: dict[str, Any], missing_documents: list[str]) -> str:
        current_status = str(application.get("status") or "draft")
        if current_status == "submitted":
            return current_status
        if missing_documents:
            return "needs_documents"
        if application.get("review_json", {}).get("readiness") == "ready" or current_status == "reviewed":
            return "reviewed"
        if current_status in {"needs_documents", "blocked_missing_documents"}:
            return "draft"
        return current_status

    def _refresh_application_compliance(self, application: dict[str, Any], *, persist: bool) -> dict[str, Any]:
        package = dict(application.get("package_json", {}) or {})
        compliance = dict(package.get("compliance", {}) or {})
        required_documents = list(compliance.get("required_documents", []))
        if not required_documents:
            return application
        available_documents, missing_documents = self._application_document_state(application["company_id"], required_documents)
        if (
            compliance.get("available_documents", []) == available_documents
            and compliance.get("missing_documents", []) == missing_documents
        ):
            return application

        compliance["available_documents"] = available_documents
        compliance["missing_documents"] = missing_documents
        package["compliance"] = compliance

        refreshed = dict(application)
        refreshed["package_json"] = package
        refreshed["status"] = self._status_after_document_refresh(application, missing_documents)
        if not persist:
            return refreshed

        return self.store.update_application(
            application["application_id"],
            {
                "status": refreshed["status"],
                "package": package,
                "review": application.get("review_json", {}),
                "exports": application.get("export_json", {}),
            },
        )

    def _build_submission_validation(self, application: dict[str, Any]) -> dict[str, Any]:
        application = self._refresh_application_compliance(application, persist=True)
        compliance = application.get("package_json", {}).get("compliance", {})
        return {
            "missing_documents": list(compliance.get("missing_documents", [])),
            "review_readiness": application.get("review_json", {}).get("readiness", "not_reviewed"),
        }

    def _grant_amount_for_application(self, application: dict[str, Any]) -> float:
        budget = application.get("package_json", {}).get("budget", {})
        try:
            return float(budget.get("grant_requested", 0.0) or 0.0)
        except (TypeError, ValueError):
            return 0.0

    def _submission_status_from_validation(self, validation: dict[str, Any], ready_status: str) -> str:
        if validation.get("missing_documents"):
            return "blocked_missing_documents"
        if validation.get("review_readiness") != "ready":
            return "blocked_review"
        return ready_status

    def _assert_submission_ready(self, session: dict[str, Any], application: dict[str, Any], *, action_name: str) -> dict[str, Any]:
        validation = self._build_submission_validation(application)
        stored_validation = session.get("validation_json", {}) or {}
        missing_documents = list(validation.get("missing_documents", []))
        review_readiness = str(validation.get("review_readiness", "not_reviewed") or "not_reviewed")
        if missing_documents:
            raise ValueError(f"Cannot {action_name} while documents are missing: {', '.join(missing_documents)}")
        if review_readiness != "ready":
            raise ValueError(f"Cannot {action_name} until application review is ready. Current readiness: {review_readiness}")
        if stored_validation != validation:
            self.store.update_submission_session(
                session["session_id"],
                {
                    "status": self._submission_status_from_validation(validation, session.get("status", "draft_prepared")),
                    "portal_url": session.get("portal_url", ""),
                    "draft_reference": session.get("draft_reference", ""),
                    "prepared_fields": session.get("prepared_fields_json", {}),
                    "validation": validation,
                    "receipt": session.get("receipt_json", {}),
                },
            )
        return validation

    def save_company_profile(self, payload: dict[str, Any], actor: str) -> dict[str, Any]:
        profile = self.store.upsert_company_profile(payload["company_id"], payload)
        self.store.record_audit_event(actor, "funding.company_profile.saved", {"company_id": payload["company_id"]}, company_id=payload["company_id"])
        return profile

    def get_company_profile(self, company_id: str) -> dict[str, Any]:
        return self.store.get_company_profile(company_id) or self._empty_company_profile(company_id)

    def get_company_registry_sync(self, company_id: str) -> dict[str, Any]:
        profile = self.store.get_company_profile(company_id) or {}
        registry_sync = profile.get("registry_sync") or {}
        return {
            "company_id": company_id,
            "available": bool(registry_sync),
            "registry_sync": registry_sync,
        }

    def sync_company_registry(self, company_id: str, krs: str, actor: str, apply_profile: bool = True) -> dict[str, Any]:
        result = fetch_krs_company_profile(krs)
        current = self.get_company_profile(company_id)
        patch = dict(result.get("profile_patch") or {})
        if apply_profile:
            merged = {**current, **{key: value for key, value in patch.items() if value not in ("", [], None)}}
            merged["company_id"] = company_id
            profile = self.store.upsert_company_profile(company_id, merged)
        else:
            profile = current
        registry_sync = result.get("registry_sync") or {}
        self.store.record_audit_event(
            actor,
            "funding.company.registry_synced",
            {
                "krs": patch.get("registration_id", krs),
                "source": registry_sync.get("source", ""),
                "applied": apply_profile,
                "financial_filings": len(registry_sync.get("financial_filings", [])),
            },
            company_id=company_id,
        )
        return {**result, "company_profile": profile, "applied": apply_profile}

    def get_company_readiness(self, company_id: str) -> dict[str, Any]:
        profile = self.get_company_profile(company_id)
        documents = self.store.list_company_documents(company_id)
        return {
            "company_id": company_id,
            "profile": profile,
            "documents": documents,
            **_company_readiness(profile, documents),
        }

    def add_company_document(self, payload: dict[str, Any], actor: str) -> dict[str, Any]:
        self._require_company(payload["company_id"])
        item = self.store.add_company_document(payload["company_id"], payload)
        for application in self.store.list_applications(payload["company_id"]):
            self._refresh_application_compliance(application, persist=True)
        self.store.record_audit_event(
            actor,
            "funding.company_document.saved",
            {"company_id": payload["company_id"], "document_type": payload.get("document_type", ""), "filename": payload.get("filename", "")},
            company_id=payload["company_id"],
        )
        return item

    def list_sources(self) -> dict[str, Any]:
        programmes = self.store.list_programmes()
        calls = self.store.list_calls()
        calls_by_programme: dict[str, int] = {}
        for call in calls:
            pid = str(call.get("programme_id") or "")
            calls_by_programme[pid] = calls_by_programme.get(pid, 0) + 1
        grouped: dict[str, dict[str, Any]] = {}
        for programme in programmes:
            source_id = str(programme.get("source_id") or "manual")
            item = grouped.setdefault(
                source_id,
                {
                    "source_id": source_id,
                    "label": "Import ręczny" if source_id == "manual" else "Skan źródeł publicznych",
                    "scan_mode": "live_manual" if source_id == "manual" else "source_snapshot",
                    "programmes": 0,
                    "calls": 0,
                    "available": True,
                },
            )
            item["programmes"] += 1
            item["calls"] += calls_by_programme.get(str(programme.get("programme_id") or ""), 0)
        grouped.setdefault(
            "manual",
            {
                "source_id": "manual",
                "label": "Import ręczny",
                "scan_mode": "live_manual",
                "programmes": 0,
                "calls": 0,
                "available": True,
            },
        )
        return {"sources": list(grouped.values())}

    def create_programme(self, payload: dict[str, Any], actor: str) -> dict[str, Any]:
        programme = self.store.create_programme(payload)
        self.store.record_audit_event(actor, "funding.programme.created", {"programme_id": programme["programme_id"], "name": programme["name"]})
        return programme

    def list_programmes(self) -> dict[str, Any]:
        return {"programmes": self.store.list_programmes()}

    def create_call(self, payload: dict[str, Any], actor: str) -> dict[str, Any]:
        if not self.store.get_programme(payload["programme_id"]):
            raise ValueError(f"Programme '{payload['programme_id']}' not found")
        call = self.store.create_call(payload)
        self.store.record_audit_event(actor, "funding.call.created", {"call_id": call["call_id"], "title": call["title"]})
        return call

    def list_calls(self) -> dict[str, Any]:
        return {"calls": self.store.list_calls()}

    def search_calls(self, payload: dict[str, Any]) -> dict[str, Any]:
        calls = self.store.list_calls()
        profile = self.store.get_company_profile(payload.get("company_id", "default"))
        now = _now()
        query_tokens = _expand_query_tokens(_clean_words([payload.get("query", ""), payload.get("theme", ""), payload.get("beneficiary_type", "")]))
        filtered = []
        company_keywords = _keywords_for_company(profile or {})
        company_budget = float(profile.get("annual_revenue", 0) or 0) if profile else 0.0
        company_sme = str(profile.get("sme_status", "")).lower() if profile else ""
        for call in calls:
            if payload.get("open_only", True) and call.get("closes_at") and call["closes_at"] < now:
                continue
            if payload.get("country") and not _country_compatible(payload["country"], call.get("country", "")):
                continue
            call_tokens = _keywords_for_call(call)
            if query_tokens and not (query_tokens & call_tokens):
                continue
            if payload.get("beneficiary_type"):
                beneficiaries = {_beneficiary_key(item) for item in call.get("target_beneficiaries_json", [])}
                if beneficiaries and _beneficiary_key(payload["beneficiary_type"]) not in beneficiaries:
                    continue
            min_budget = payload.get("budget_min")
            max_budget = payload.get("budget_max")
            call_min = float(call.get("min_project_budget", 0) or 0)
            call_max = float(call.get("max_project_budget", 0) or 0)
            if min_budget is not None and call_max and call_max < min_budget:
                continue
            if max_budget is not None and call_min and call_min > max_budget:
                continue
            fit_signals = []
            if company_keywords:
                overlap = _overlap_score(company_keywords, call_tokens)
                fit_signals.append(overlap)
            if company_budget and call_min <= company_budget * 4:
                fit_signals.append(80.0)
            if company_sme and any(company_sme in str(item).lower() for item in call.get("target_beneficiaries_json", [])):
                fit_signals.append(90.0)
            filtered.append({
                **call,
                "fit_hint": round(sum(fit_signals) / max(len(fit_signals), 1), 2) if fit_signals else 0.0,
            })
        filtered.sort(key=lambda item: (item.get("fit_hint", 0), -(item.get("closes_at") or 9e18)), reverse=True)
        return {"calls": filtered}

    def generate_ideas(self, company_id: str, limit: int = 5) -> dict[str, Any]:
        profile = self._require_company(company_id)
        calls = self.store.list_calls()
        now = _now()
        active_calls = [call for call in calls if not call.get("closes_at") or float(call.get("closes_at") or 0) >= now]
        ideas: list[dict[str, Any]] = []
        readiness = self.get_company_readiness(company_id)
        company_keywords = _keywords_for_company(profile)
        if active_calls:
            ranked_calls = []
            for call in active_calls:
                overlap = _overlap_score(company_keywords, _keywords_for_call(call))
                ranked_calls.append((overlap, call))
            ranked_calls.sort(key=lambda item: item[0], reverse=True)
            for overlap, call in ranked_calls[: max(limit, 1)]:
                category = _categorize(profile, call)
                grant_estimate = round((float(call.get("grant_intensity_pct", 0) or 0) / 100.0) * float(call.get("max_project_budget", 0) or 0), 2)
                risk_level = "low" if readiness["readiness_score"] >= 75 else "medium" if readiness["readiness_score"] >= 50 else "high"
                chance_pct = _clamp((overlap * 0.45) + (readiness["readiness_score"] * 0.55))
                recommendation = "aplikowac" if chance_pct >= 70 else "przygotowac" if chance_pct >= 50 else "obserwowac"
                themes = ", ".join(call.get("themes_json", [])[:3]) or call.get("title", "")
                ideas.append(
                    {
                        "title": f"{category.title()} dla {profile.get('legal_name', 'firmy')}",
                        "category": category,
                        "problem": f"Projekt odpowiada na priorytety naboru: {themes}.",
                        "solution": f"Wykorzystanie zasobow firmy do przygotowania wdrozenia zgodnego z zakresem {call.get('title', '')}.",
                        "objective": f"Zweryfikowac i wdrozyc rozwiazanie zgodne z celami programu {call.get('code', '') or call.get('title', '')}.",
                        "recommended_call_id": call["call_id"],
                        "recommended_programme_id": call["programme_id"],
                        "recommended_country": call.get("country", ""),
                        "recommended_institution": call.get("metadata_json", {}).get("institution", call.get("region", "")) or call.get("region", ""),
                        "budget_estimate": float(call.get("max_project_budget", 0) or 0) or 250000.0,
                        "grant_estimate": grant_estimate,
                        "required_partner_types": call.get("required_partner_types_json", []),
                        "difficulty": "medium" if chance_pct >= 50 else "high",
                        "risk_level": risk_level,
                        "chance_pct": chance_pct,
                        "recommendation": recommendation,
                        "rationale": f"Overlap tematyczny: {overlap}%. Gotowosc firmy: {readiness['readiness_score']}%.",
                    }
                )
        else:
            category = _categorize(profile)
            ideas.append(
                {
                    "title": f"{category.title()} dla {profile.get('legal_name', 'firmy')}",
                    "category": category,
                    "problem": "Brak zaimportowanych naborow. System zbudowal pomysl na podstawie profilu firmy.",
                    "solution": "Przygotowac koncepcje projektu i nastepnie dopasowac ja do naborow po imporcie wiedzy o programach.",
                    "objective": "Zwiekszyc gotowosc firmy do szybkiego wejscia w odpowiedni nabor.",
                    "recommended_call_id": "",
                    "recommended_country": profile.get("country", ""),
                    "recommended_institution": "",
                    "budget_estimate": 250000.0,
                    "grant_estimate": 125000.0,
                    "required_partner_types": [],
                    "difficulty": "medium",
                    "risk_level": "medium",
                    "chance_pct": _clamp(readiness["readiness_score"] * 0.7),
                    "recommendation": "przygotowac",
                    "rationale": "Dodaj przynajmniej jeden realny nabor, aby przejsc od pomyslu do dopasowania programu i scoringu.",
                }
            )
        return {"ideas": self.store.replace_ideas(company_id, ideas)}

    def get_idea(self, idea_id: str) -> dict[str, Any]:
        idea = self.store.get_idea(idea_id)
        if not idea:
            raise ValueError(f"Funding idea '{idea_id}' not found")
        return idea

    def convert_idea_to_project(self, idea_id: str, company_id: str, call_id: str | None, target_trl: int, actor: str) -> dict[str, Any]:
        idea = self.get_idea(idea_id)
        profile = self._require_company(company_id)
        chosen_call_id = call_id or idea.get("recommended_call_id", "")
        for existing in self.store.list_projects(company_id):
            if existing.get("idea_id") == idea_id and (existing.get("call_id") or "") == (chosen_call_id or ""):
                self.store.record_audit_event(
                    actor,
                    "funding.idea.converted_to_existing_project",
                    {"idea_id": idea_id, "project_id": existing["project_id"]},
                    company_id=company_id,
                )
                return {"project": existing, "reused_existing": True}
        project = self.store.create_project(
            {
                "company_id": company_id,
                "idea_id": idea_id,
                "call_id": chosen_call_id,
                "title": idea["title"],
                "status": "draft",
                "summary": idea["solution"],
                "objective": idea["objective"],
                "category": idea["category"],
                "budget_total": idea.get("budget_estimate", 0),
                "grant_requested": idea.get("grant_estimate", 0),
                "trl": target_trl,
                "target_markets": profile.get("export_markets", []),
                "partner_needs": idea.get("required_partner_types", []),
                "work_packages": ["Analysis", "Build", "Validation", "Commercialisation"],
                "milestones": ["Project kickoff", "Prototype ready", "Validation complete", "Application package ready"],
                "risk_register": [idea.get("risk_level", "medium")],
            }
        )
        self.store.record_audit_event(actor, "funding.idea.converted_to_project", {"idea_id": idea_id, "project_id": project["project_id"]}, company_id=company_id)
        return {"project": project}

    def create_project(self, payload: dict[str, Any], actor: str) -> dict[str, Any]:
        self._require_company(payload["company_id"])
        project = self.store.create_project({**payload, "status": "draft"})
        self.store.record_audit_event(actor, "funding.project.created", {"project_id": project["project_id"]}, company_id=payload["company_id"])
        return {"project": project}

    def list_projects(self, company_id: str | None = None) -> dict[str, Any]:
        return {"projects": self.store.list_projects(company_id)}

    def get_project(self, project_id: str) -> dict[str, Any]:
        return self._require_project(project_id)

    def _score_project_against_call(self, profile: dict[str, Any], documents: list[dict[str, Any]], project: dict[str, Any], call: dict[str, Any]) -> dict[str, Any]:
        company_country = profile.get("country", "")
        call_country = call.get("country", "")
        company_sme = _beneficiary_key(profile.get("sme_status", ""))
        beneficiaries = {_beneficiary_key(item) for item in call.get("target_beneficiaries_json", [])}
        project_budget = float(project.get("budget_total", 0) or 0)
        project_trl = int(project.get("trl", 0) or 0)
        required_docs = set(call.get("required_documents_json", []))
        available_docs = {item.get("document_type", "") for item in documents if item.get("status") == "available"}

        formal_components = []
        formal_reasons = []
        if _country_compatible(company_country, call_country):
            formal_components.append(100.0)
            formal_reasons.append("Kraj firmy jest kwalifikowalny.")
        else:
            formal_components.append(10.0)
            formal_reasons.append("Kraj firmy nie zgadza sie z krajem naboru.")
        if not beneficiaries or company_sme in beneficiaries:
            formal_components.append(100.0)
            formal_reasons.append("Typ beneficjenta pasuje do naboru.")
        else:
            formal_components.append(25.0)
            formal_reasons.append("Typ beneficjenta jest slabiej dopasowany.")
        min_budget = float(call.get("min_project_budget", 0) or 0)
        max_budget = float(call.get("max_project_budget", 0) or 0)
        if (min_budget == 0 or project_budget >= min_budget) and (max_budget == 0 or project_budget <= max_budget):
            formal_components.append(100.0)
            formal_reasons.append("Budzet miesci sie w zakresie naboru.")
        else:
            formal_components.append(20.0)
            formal_reasons.append("Budzet wymaga korekty do zakresu naboru.")
        if project_trl >= int(call.get("trl_min", 0) or 0) and project_trl <= int(call.get("trl_max", 9) or 9):
            formal_components.append(100.0)
            formal_reasons.append("Poziom TRL jest zgodny.")
        else:
            formal_components.append(30.0)
            formal_reasons.append("Poziom TRL nie jest w pelni zgodny.")
        closes_at = call.get("closes_at")
        if closes_at and float(closes_at) < _now():
            formal_components.append(0.0)
            formal_reasons.append("Nabor jest juz zamkniety.")
        else:
            formal_components.append(100.0)
            formal_reasons.append("Termin naboru pozwala na przygotowanie aplikacji.")
        missing_docs = sorted(required_docs - available_docs)
        formal_eligibility = round(sum(formal_components) / max(len(formal_components), 1), 2)

        project_keywords = _keywords_for_project(project) | _keywords_for_company(profile)
        call_keywords = _keywords_for_call(call)
        strategic_fit = _overlap_score(project_keywords, call_keywords) or 15.0
        innovation_score = _clamp(35 + (len(profile.get("technologies", [])) * 6) + (len(profile.get("patents", [])) * 8) + (10 if "B+R" in str(project.get("category", "")) else 0))
        market_potential = _clamp(40 + (len(profile.get("export_markets", [])) * 8) + (len(profile.get("products", [])) * 5))
        team_capacity = _clamp(30 + min(int(profile.get("employees", 0) or 0), 40) * 1.5 + len(profile.get("team_competencies", [])) * 4)
        own_contribution = max(project_budget - float(project.get("grant_requested", 0) or 0), 0.0)
        financial_capacity = _clamp(20 + (float(profile.get("annual_revenue", 0) or 0) / max(own_contribution or 1.0, 1.0)) * 15)
        consortium_strength = 100.0
        partner_needs = set(project.get("partner_needs", []))
        if call.get("requires_consortium") or call.get("required_partner_types_json"):
            required_partner_types = set(call.get("required_partner_types_json", []))
            missing_partner_types = sorted(required_partner_types - partner_needs)
            consortium_strength = _clamp(100 - (len(missing_partner_types) * 25))
        else:
            missing_partner_types = []
        documentation_readiness = _clamp(100 - (len(missing_docs) * 20))
        risk_penalties = 0.0
        risks = []
        if missing_docs:
            risk_penalties += min(25, len(missing_docs) * 5)
            risks.append(f"Brakuje dokumentow: {', '.join(missing_docs)}")
        if missing_partner_types:
            risk_penalties += min(20, len(missing_partner_types) * 6)
            risks.append(f"Brakuje partnerow: {', '.join(missing_partner_types)}")
        if closes_at:
            days_left = (float(closes_at) - _now()) / 86400.0
            if days_left < 0:
                risk_penalties += 35
                risks.append("Nabor jest zamkniety, nie powinien byc rekomendowany do nowego wniosku.")
            elif days_left < 14:
                risk_penalties += 10
                risks.append("Termin naboru jest blisko.")
        fit_score = _clamp((formal_eligibility * 0.5) + (strategic_fit * 0.5))
        probability = _clamp(
            (formal_eligibility * 0.20)
            + (strategic_fit * 0.15)
            + (innovation_score * 0.15)
            + (market_potential * 0.10)
            + (team_capacity * 0.10)
            + (financial_capacity * 0.10)
            + (consortium_strength * 0.10)
            + (documentation_readiness * 0.10)
            - risk_penalties
        )
        improvements = []
        if missing_docs:
            improvements.append({"action": "Uzupelnij dokumenty", "impact_probability_pct": min(20, len(missing_docs) * 6)})
        if missing_partner_types:
            improvements.append({"action": "Dodaj brakujacych partnerow", "impact_probability_pct": min(15, len(missing_partner_types) * 5)})
        if financial_capacity < 60:
            improvements.append({"action": "Zmniejsz budzet lub zwieksz wklad wlasny", "impact_probability_pct": 8})
        strengths = [
            *formal_reasons[:2],
            f"Dopasowanie tematyczne: {strategic_fit}%",
            f"Gotowosc dokumentacyjna: {documentation_readiness}%",
        ]
        programme = self.store.get_programme(call["programme_id"]) if call.get("programme_id") else None
        return {
            "call_id": call["call_id"],
            "project_id": project["project_id"],
            "fit_score": fit_score,
            "success_probability": probability,
            "readiness_score": documentation_readiness,
            "risk_score": _clamp(risk_penalties),
            "formal_eligibility": formal_eligibility,
            "strategic_fit": strategic_fit,
            "innovation_score": innovation_score,
            "market_potential": market_potential,
            "team_capacity": team_capacity,
            "financial_capacity": financial_capacity,
            "consortium_strength": consortium_strength,
            "documentation_readiness": documentation_readiness,
            "risk_penalties": risk_penalties,
            "strengths": strengths,
            "risks": risks,
            "improvements": improvements,
            "missing_documents": missing_docs,
            "missing_partner_types": missing_partner_types,
            "confidence": _clamp((fit_score + documentation_readiness) / 2.0),
            "evidence": {
                "programme_name": programme.get("name", "") if programme else "",
                "call_title": call.get("title", ""),
                "portal_url": call.get("portal_url", ""),
                "evaluated_at": _now(),
                "company_fields_used": ["country", "sme_status", "annual_revenue", "technologies", "products", "team_competencies"],
            },
        }

    def run_matching(self, project_id: str, call_id: str | None = None, top_k: int = 5) -> dict[str, Any]:
        project = self._require_project(project_id)
        profile = self._require_company(project["company_id"])
        documents = self.store.list_company_documents(project["company_id"])
        if call_id:
            calls = [self._require_call(call_id)]
        elif project.get("call_id"):
            calls = [self._require_call(project["call_id"])]
        else:
            calls = self.store.list_calls()
        if not calls:
            raise ValueError("No funding calls available. Import or create at least one call first.")
        matches = [self._score_project_against_call(profile, documents, project, call) for call in calls]
        matches.sort(key=lambda item: (item["success_probability"], item["fit_score"]), reverse=True)
        top_matches = matches[: max(top_k, 1)]
        self.store.replace_matches(project_id, top_matches)
        if top_matches:
            best = top_matches[0]
            project["status"] = "matched"
            project["best_match_call_id"] = best["call_id"]
            project["best_match_probability"] = best["success_probability"]
            self.store.update_project(project_id, project)
        return {"project_id": project_id, "matches": self.store.list_matches(project_id)}

    def get_matching_results(self, project_id: str) -> dict[str, Any]:
        self._require_project(project_id)
        return {"project_id": project_id, "matches": self.store.list_matches(project_id)}

    def check_eligibility(self, project_id: str, call_id: str | None = None) -> dict[str, Any]:
        project = self._require_project(project_id)
        profile = self._require_company(project["company_id"])
        documents = self.store.list_company_documents(project["company_id"])
        selected_call = self._require_call(call_id or project.get("call_id") or self.store.list_calls()[0]["call_id"])
        match = self._score_project_against_call(profile, documents, project, selected_call)
        return {
            "project_id": project_id,
            "call_id": selected_call["call_id"],
            "eligible": match["formal_eligibility"] >= 60,
            "formal_eligibility": match["formal_eligibility"],
            "missing_documents": match["missing_documents"],
            "missing_partner_types": match["missing_partner_types"],
            "reasons": match["strengths"] + match["risks"],
        }

    def run_scoring(self, project_id: str, call_id: str | None = None) -> dict[str, Any]:
        result = self.run_matching(project_id, call_id=call_id, top_k=1)
        best_match = result["matches"][0]
        return {
            "project_id": project_id,
            "call_id": best_match["call_id"],
            "grant_fit_score": best_match["fit_score"],
            "grant_success_probability": best_match["success_probability"],
            "confidence": best_match["confidence"],
            "strengths": best_match["strengths"],
            "risks": best_match["risks"],
            "improvements": best_match["improvements"],
            "simulation": [
                {
                    "scenario": item["action"],
                    "projected_probability": _clamp(best_match["success_probability"] + item["impact_probability_pct"]),
                }
                for item in best_match["improvements"]
            ],
        }

    def get_scoring(self, project_id: str) -> dict[str, Any]:
        matches = self.store.list_matches(project_id)
        if not matches:
            return self.run_scoring(project_id)
        best_match = matches[0]
        return {
            "project_id": project_id,
            "call_id": best_match["call_id"],
            "grant_fit_score": best_match["fit_score"],
            "grant_success_probability": best_match["success_probability"],
            "confidence": best_match["confidence"],
            "strengths": best_match["strengths"],
            "risks": best_match["risks"],
            "improvements": best_match["improvements"],
        }

    def analyze_consortium(self, project_id: str) -> dict[str, Any]:
        project = self._require_project(project_id)
        matches = self.store.list_matches(project_id)
        if not matches:
            matches = self.run_matching(project_id, top_k=1)["matches"]
        best = matches[0]
        required = list(best.get("missing_partner_types", []))
        if not required and project.get("partner_needs"):
            required = list(project.get("partner_needs", []))
        role_map = [
            {"role": partner_type, "responsibility": f"Wnosi kompetencje typu {partner_type} do konsorcjum."}
            for partner_type in required
        ]
        return {
            "project_id": project_id,
            "required_partner_types": required,
            "recommended_roles": role_map,
            "consortium_strength": best.get("consortium_strength", 100.0),
            "readiness": "needs_partners" if required else "self_sufficient",
        }

    def search_partners(self, payload: dict[str, Any]) -> dict[str, Any]:
        project = self._require_project(payload["project_id"])
        company_id = payload.get("company_id", project["company_id"])
        if payload.get("candidates"):
            persisted = []
            for candidate in payload["candidates"]:
                persisted.append(
                    {
                        "name": candidate["name"],
                        "partner_type": candidate["partner_type"],
                        "country": candidate.get("country", ""),
                        "expertise": candidate.get("expertise", []),
                        "grant_track_record": candidate.get("grant_track_record", 0),
                        "contact_email": candidate.get("contact_email", ""),
                        "metadata": candidate.get("metadata", {}),
                    }
                )
            self.store.replace_partner_candidates(project["project_id"], company_id, persisted)
        candidates = self.store.list_partner_candidates(project["project_id"])
        query_tokens = _clean_words([payload.get("query", ""), payload.get("partner_type", "")])
        project_keywords = _keywords_for_project(project)
        results = []
        for candidate in candidates:
            expertise_tokens = _clean_words(candidate.get("expertise_json", []))
            if payload.get("partner_type") and payload["partner_type"].strip().lower() != str(candidate.get("partner_type", "")).lower():
                continue
            if query_tokens and not (query_tokens & (expertise_tokens | _clean_words([candidate.get("name", "")]))):
                continue
            candidate["score"] = _clamp((_overlap_score(project_keywords, expertise_tokens) * 0.7) + (float(candidate.get("grant_track_record", 0) or 0) * 6))
            results.append(candidate)
        results.sort(key=lambda item: item.get("score", 0), reverse=True)
        self.store.replace_partner_candidates(project["project_id"], company_id, results)
        return {"project_id": project["project_id"], "partners": self.store.list_partner_candidates(project["project_id"])}

    def shortlist_partners(self, project_id: str, limit: int = 5) -> dict[str, Any]:
        project = self._require_project(project_id)
        partners = self.store.list_partner_candidates(project_id)[: max(limit, 1)]
        return {"project_id": project["project_id"], "shortlist": partners}

    def generate_outreach(self, project_id: str, partner_ids: list[str]) -> dict[str, Any]:
        project = self._require_project(project_id)
        partners = [item for item in self.store.list_partner_candidates(project_id) if item["partner_id"] in set(partner_ids)]
        messages = []
        for partner in partners:
            messages.append(
                {
                    "partner_id": partner["partner_id"],
                    "message_type": "intro_email",
                    "subject": f"Propozycja wspolnego projektu grantowego: {project['title']}",
                    "body": (
                        f"Dzien dobry,\n\n"
                        f"przygotowujemy projekt '{project['title']}' i szukamy partnera typu {partner.get('partner_type', '')}. "
                        f"Panstwa profil kompetencyjny ({', '.join(partner.get('expertise_json', []))}) wyglada na dobrze dopasowany. "
                        f"Czy mozemy omowic mozliwa wspolprace i zakres roli w konsorcjum?\n\n"
                        "Pozdrawiamy,\nAEIS Funding Autopilot"
                    ),
                }
            )
        return {"project_id": project_id, "messages": self.store.replace_outreach_messages(project_id, messages)}

    def _build_application_package(self, profile: dict[str, Any], project: dict[str, Any], call: dict[str, Any], match: dict[str, Any]) -> dict[str, Any]:
        required_documents = list(call.get("required_documents_json", []))
        available_documents, missing_documents = self._application_document_state(project["company_id"], required_documents)
        budget_total = float(project.get("budget_total", 0) or 0)
        grant_requested = float(project.get("grant_requested", 0) or 0)
        own_contribution = max(budget_total - grant_requested, 0.0)
        return {
            "executive_summary": (
                f"Projekt '{project['title']}' odpowiada na nabor '{call.get('title', '')}'. "
                f"Grant Fit Score: {match['fit_score']}%, przewidywana szansa sukcesu: {match['success_probability']}%."
            ),
            "project_description": {
                "problem": project.get("summary", ""),
                "objective": project.get("objective", ""),
                "category": project.get("category", ""),
                "target_markets": project.get("target_markets", []),
            },
            "innovation_case": {
                "innovation_score": match["innovation_score"],
                "technologies": profile.get("technologies", []),
                "patents": profile.get("patents", []),
            },
            "market_case": {
                "market_potential": match["market_potential"],
                "products": profile.get("products", []),
                "export_markets": profile.get("export_markets", []),
            },
            "implementation_plan": {
                "work_packages": project.get("work_packages", []),
                "milestones": project.get("milestones", []),
                "risks": match["risks"],
            },
            "budget": {
                "budget_total": budget_total,
                "grant_requested": grant_requested,
                "own_contribution": own_contribution,
                "grant_intensity_pct": call.get("grant_intensity_pct", 0),
                "cost_categories": call.get("eligible_costs_json", []),
            },
            "consortium": {
                "required_partner_types": call.get("required_partner_types_json", []),
                "current_partner_needs": project.get("partner_needs", []),
            },
            "compliance": {
                "required_documents": required_documents,
                "available_documents": available_documents,
                "missing_documents": missing_documents,
                "eligibility_summary": match["strengths"],
            },
        }

    def create_application(self, project_id: str, company_id: str, call_id: str | None = None, actor: str = "workspace-default") -> dict[str, Any]:
        project = self._require_project(project_id)
        profile = self._require_company(company_id)
        matches = self.store.list_matches(project_id) or self.run_matching(project_id, call_id=call_id, top_k=1)["matches"]
        best = matches[0]
        selected_call = self._require_call(call_id or best["call_id"])
        package = self._build_application_package(profile, project, selected_call, best)
        status = "draft"
        if package["compliance"]["missing_documents"]:
            status = "needs_documents"
        application = self.store.create_application(
            {
                "company_id": company_id,
                "project_id": project_id,
                "call_id": selected_call["call_id"],
                "status": status,
                "package": package,
                "review": {},
                "exports": {},
            }
        )
        self.store.record_audit_event(actor, "funding.application.created", {"application_id": application["application_id"], "project_id": project_id}, company_id=company_id, application_id=application["application_id"])
        return application

    def get_application(self, application_id: str) -> dict[str, Any]:
        return self._refresh_application_compliance(self._require_application(application_id), persist=False)

    def get_application_documents(self, application_id: str) -> dict[str, Any]:
        application = self._refresh_application_compliance(self._require_application(application_id), persist=False)
        package = application.get("package_json", {})
        compliance = package.get("compliance", {})
        return {
            "application_id": application_id,
            "required_documents": compliance.get("required_documents", []),
            "available_documents": compliance.get("available_documents", []),
            "missing_documents": compliance.get("missing_documents", []),
        }

    def review_application(self, application_id: str, review_modes: list[str]) -> dict[str, Any]:
        application = self._refresh_application_compliance(self._require_application(application_id), persist=True)
        package = application.get("package_json", {})
        findings = []
        compliance = package.get("compliance", {})
        if "formal" in review_modes and compliance.get("missing_documents"):
            findings.append({"reviewer": "formal", "severity": "high", "message": f"Missing documents: {', '.join(compliance['missing_documents'])}"})
        if "financial" in review_modes and float(package.get("budget", {}).get("own_contribution", 0) or 0) <= 0:
            findings.append({"reviewer": "financial", "severity": "medium", "message": "Own contribution is zero or not defined."})
        if "technical" in review_modes and len(package.get("implementation_plan", {}).get("work_packages", [])) < 3:
            findings.append({"reviewer": "technical", "severity": "medium", "message": "Implementation plan is too shallow."})
        if "market" in review_modes and not package.get("market_case", {}).get("export_markets"):
            findings.append({"reviewer": "market", "severity": "low", "message": "No target markets defined."})
        readiness = "ready" if not any(item["severity"] == "high" for item in findings) else "blocked"
        review = {"review_modes": review_modes, "findings": findings, "readiness": readiness, "reviewed_at": _now()}
        updated = self.store.update_application(
            application_id,
            {
                "status": "reviewed" if readiness == "ready" else "needs_documents",
                "package": package,
                "review": review,
                "exports": application.get("export_json", {}),
            },
        )
        return {"application": updated, "review": review}

    def export_application(self, application_id: str) -> dict[str, Any]:
        application = self._refresh_application_compliance(self._require_application(application_id), persist=True)
        export_dir = funding_results_root() / application_id
        export_dir.mkdir(parents=True, exist_ok=True)
        package = application.get("package_json", {})
        review = application.get("review_json", {})

        json_path = export_dir / "application_package.json"
        json_path.write_text(json.dumps(package, ensure_ascii=True, indent=2), encoding="utf-8")

        md_path = export_dir / "application_summary.md"
        md_lines = [
            "# Wniosek o finansowanie",
            "",
            "## Podsumowanie wykonawcze",
            "",
            package.get("executive_summary", ""),
            "",
        ]
        for section_name, section_value in package.items():
            if section_name == "executive_summary":
                continue
            md_lines.extend(
                [
                    f"## {_export_section_label(section_name)}",
                    json.dumps(section_value, ensure_ascii=False, indent=2)
                    if isinstance(section_value, (dict, list))
                    else str(section_value),
                    "",
                ]
            )
        md_path.write_text("\n".join(md_lines), encoding="utf-8")

        csv_path = export_dir / "budget.csv"
        with csv_path.open("w", encoding="utf-8", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(["metryka", "wartosc"])
            for key, value in package.get("budget", {}).items():
                writer.writerow([_export_budget_label(key), value])

        review_path = export_dir / "review.json"
        review_path.write_text(json.dumps(review, ensure_ascii=True, indent=2), encoding="utf-8")

        generated_files = {
            "json": str(json_path),
            "markdown": str(md_path),
            "csv": str(csv_path),
            "review": str(review_path),
        }

        try:
            from docx import Document  # type: ignore

            doc = Document()
            doc.add_heading("Wniosek o finansowanie", level=1)
            doc.add_heading("Podsumowanie wykonawcze", level=2)
            doc.add_paragraph(package.get("executive_summary", ""))
            for section_name, section_value in package.items():
                if section_name == "executive_summary":
                    continue
                doc.add_heading(_export_section_label(section_name), level=2)
                doc.add_paragraph(json.dumps(section_value, ensure_ascii=False, indent=2) if isinstance(section_value, (dict, list)) else str(section_value))
            docx_path = export_dir / "application.docx"
            doc.save(docx_path)
            generated_files["docx"] = str(docx_path)
        except Exception:
            pass

        try:
            from openpyxl import Workbook  # type: ignore

            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Budzet"
            sheet.append(["Metryka", "Wartosc"])
            for key, value in package.get("budget", {}).items():
                sheet.append([_export_budget_label(key), value])
            xlsx_path = export_dir / "budget.xlsx"
            workbook.save(xlsx_path)
            generated_files["xlsx"] = str(xlsx_path)
        except Exception:
            xlsx_path = export_dir / "budget.xlsx"
            _write_budget_xlsx(xlsx_path, package.get("budget", {}))
            generated_files["xlsx"] = str(xlsx_path)

        try:
            from reportlab.lib.pagesizes import A4  # type: ignore
            from reportlab.pdfgen import canvas  # type: ignore

            pdf_path = export_dir / "application.pdf"
            pdf = canvas.Canvas(str(pdf_path), pagesize=A4)
            text = pdf.beginText(40, 800)
            text.textLine("Wniosek o finansowanie")
            text.textLine("Podsumowanie wykonawcze")
            text.textLine(package.get("executive_summary", ""))
            text.textLine("Wyniki przegladu")
            for finding in review.get("findings", []):
                text.textLine(f"- {finding.get('reviewer')}: {finding.get('message')}")
            pdf.drawText(text)
            pdf.save()
            generated_files["pdf"] = str(pdf_path)
        except Exception:
            pdf_path = export_dir / "application.pdf"
            pdf_lines = [
                "Wniosek o finansowanie",
                "Podsumowanie wykonawcze",
                package.get("executive_summary", ""),
                "Wyniki przegladu",
                *[
                    f"- {finding.get('reviewer')}: {finding.get('message')}"
                    for finding in review.get("findings", [])
                ],
            ]
            _write_text_pdf(pdf_path, pdf_lines)
            generated_files["pdf"] = str(pdf_path)

        zip_path = export_dir / "application_bundle.zip"
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for path in generated_files.values():
                archive.write(path, Path(path).name)
        generated_files["zip"] = str(zip_path)

        updated = self.store.update_application(
            application_id,
            {
                "status": application.get("status", "draft"),
                "package": package,
                "review": review,
                "exports": generated_files,
            },
        )
        return {"application": updated, "exports": generated_files}

    def get_application_export_file(self, application_id: str, artifact_type: str) -> Path:
        artifact = artifact_type.strip().lower()
        if artifact not in EXPORT_ARTIFACT_MEDIA_TYPES:
            raise ValueError(f"Unsupported export artifact: {artifact_type}")
        application = self._require_application(application_id)
        exports = dict(application.get("export_json", {}) or {})
        path_value = exports.get(artifact)
        if not path_value:
            exports = self.export_application(application_id)["exports"]
            path_value = exports.get(artifact)
        if not path_value:
            raise ValueError(f"Export artifact is not available: {artifact}")

        export_path = Path(str(path_value)).resolve()
        export_root = funding_results_root().resolve()
        try:
            export_path.relative_to(export_root)
        except ValueError as exc:
            raise ValueError("Export artifact is outside the funding results root") from exc
        if not export_path.is_file():
            raise ValueError(f"Export artifact is missing on disk: {artifact}")
        return export_path

    def prepare_submission(self, application_id: str, portal_url: str, actor: str) -> dict[str, Any]:
        application = self._refresh_application_compliance(self._require_application(application_id), persist=True)
        validation = self._build_submission_validation(application)
        status = self._submission_status_from_validation(validation, "draft_prepared")
        session = self.store.create_submission_session(
            {
                "application_id": application_id,
                "status": status,
                "portal_url": portal_url,
                "prepared_fields": {},
                "validation": validation,
                "receipt": {},
            }
        )
        self.store.record_audit_event(actor, "funding.submission.prepared", {"application_id": application_id, "session_id": session["session_id"]}, company_id=application["company_id"], application_id=application_id)
        return session

    def fill_submission(self, session_id: str, actor: str) -> dict[str, Any]:
        session = self._require_submission_session(session_id)
        application = self._refresh_application_compliance(self._require_application(session["application_id"]), persist=True)
        package = application.get("package_json", {})
        validation = self._build_submission_validation(application)
        prepared_fields = {
            "company_name": self._require_company(application["company_id"]).get("legal_name", ""),
            "project_title": self._require_project(application["project_id"]).get("title", ""),
            "project_summary": package.get("executive_summary", ""),
            "budget_total": package.get("budget", {}).get("budget_total", 0),
            "grant_requested": package.get("budget", {}).get("grant_requested", 0),
            "required_documents": package.get("compliance", {}).get("required_documents", []),
        }
        updated = self.store.update_submission_session(
            session_id,
            {
                "status": self._submission_status_from_validation(validation, "form_mapping_ready"),
                "portal_url": session.get("portal_url", ""),
                "draft_reference": session.get("draft_reference", ""),
                "prepared_fields": prepared_fields,
                "validation": validation,
                "receipt": session.get("receipt_json", {}),
            },
        )
        self.store.record_audit_event(actor, "funding.submission.filled", {"session_id": session_id}, company_id=application["company_id"], application_id=application["application_id"])
        return updated

    def save_draft(self, session_id: str, actor: str) -> dict[str, Any]:
        session = self._require_submission_session(session_id)
        application = self._refresh_application_compliance(self._require_application(session["application_id"]), persist=True)
        validation = self._build_submission_validation(application)
        draft_reference = session.get("draft_reference") or f"draft-{session_id}"
        updated = self.store.update_submission_session(
            session_id,
            {
                "status": self._submission_status_from_validation(validation, "draft_saved"),
                "portal_url": session.get("portal_url", ""),
                "draft_reference": draft_reference,
                "prepared_fields": session.get("prepared_fields_json", {}),
                "validation": validation,
                "receipt": session.get("receipt_json", {}),
            },
        )
        self.store.record_audit_event(actor, "funding.submission.draft_saved", {"session_id": session_id, "draft_reference": draft_reference}, company_id=application["company_id"], application_id=application["application_id"])
        return updated

    def request_approval(self, session_id: str, action_type: str, requested_by: str, notes: str) -> dict[str, Any]:
        session = self._require_submission_session(session_id)
        application = self._refresh_application_compliance(self._require_application(session["application_id"]), persist=True)
        validation = self._assert_submission_ready(session, application, action_name="request approval")
        governance_ticket_id = submit_submission_ticket(
            application_id=application["application_id"],
            session_id=session_id,
            portal=session.get("portal_url", ""),
            amount=self._grant_amount_for_application(application),
        )
        event = self.store.create_approval_event(
            {
                "application_id": application["application_id"],
                "session_id": session_id,
                "action_type": action_type,
                "status": "pending",
                "requested_by": requested_by,
                "payload": {
                    "notes": notes,
                    "application_version": application.get("updated_at"),
                    "validation": validation,
                    "governance_ticket_id": governance_ticket_id,
                    "human_gate_state": "pending",
                },
            }
        )
        self.store.update_submission_session(
            session_id,
            {
                "status": "awaiting_approval",
                "portal_url": session.get("portal_url", ""),
                "draft_reference": session.get("draft_reference", ""),
                "prepared_fields": session.get("prepared_fields_json", {}),
                "validation": validation,
                "receipt": session.get("receipt_json", {}),
            },
        )
        self.store.record_audit_event(requested_by, "funding.submission.approval_requested", {"session_id": session_id, "approval_event_id": event["approval_event_id"], "governance_ticket_id": governance_ticket_id}, company_id=application["company_id"], application_id=application["application_id"])
        return event

    def _require_submission_human_gate_approval(
        self,
        session: dict[str, Any],
        application: dict[str, Any],
        latest_approval: dict[str, Any],
    ) -> str:
        payload = latest_approval.get("payload_json", {}) or {}
        governance_ticket_id = str(payload.get("governance_ticket_id") or "")
        if not governance_ticket_id:
            governance_ticket_id = submit_submission_ticket(
                application_id=application["application_id"],
                session_id=session["session_id"],
                portal=session.get("portal_url", ""),
                amount=self._grant_amount_for_application(application),
            )
            self.store.update_approval_event(
                latest_approval["approval_event_id"],
                {
                    "payload": {
                        "governance_ticket_id": governance_ticket_id,
                        "human_gate_state": "pending",
                    },
                },
            )
        if not check_approved(governance_ticket_id):
            raise ValueError(
                "Human Gate approval is required before final funding submit. "
                f"Approve governance ticket {governance_ticket_id} first."
            )
        return governance_ticket_id

    def submit(self, session_id: str, approved_by: str, confirm_legal: bool, confirm_budget: bool, confirm_documents: bool, portal_submission_reference: str) -> dict[str, Any]:
        session = self._require_submission_session(session_id)
        application = self._refresh_application_compliance(self._require_application(session["application_id"]), persist=True)
        validation = self._assert_submission_ready(session, application, action_name="submit")
        approvals = self.store.list_approval_events(session_id=session_id)
        if not approvals:
            raise ValueError("Approval request is missing. Request approval before final submit.")
        if not (confirm_legal and confirm_budget and confirm_documents):
            raise ValueError("All legal, budget, and document confirmations are required.")
        if not portal_submission_reference.strip():
            raise ValueError("portal_submission_reference is required for an honest recorded submit.")
        latest = approvals[0]
        if latest.get("status") != "pending":
            raise ValueError("Latest approval request is not pending. Request a fresh approval before final submit.")
        governance_ticket_id = self._require_submission_human_gate_approval(session, application, latest)
        self.store.update_approval_event(
            latest["approval_event_id"],
            {
                "status": "approved",
                "requested_by": latest.get("requested_by", ""),
                "approved_by": approved_by,
                "approved_at": _now(),
                "payload": {
                    **latest.get("payload_json", {}),
                    "governance_ticket_id": governance_ticket_id,
                    "human_gate_state": "approved",
                    "portal_submission_reference": portal_submission_reference,
                },
            },
        )
        receipt = {
            "session_id": session_id,
            "application_id": application["application_id"],
            "submitted_at": _now(),
            "submitted_by": approved_by,
            "portal_submission_reference": portal_submission_reference,
            "draft_reference": session.get("draft_reference", ""),
        }
        updated_session = self.store.update_submission_session(
            session_id,
            {
                "status": "submitted",
                "portal_url": session.get("portal_url", ""),
                "draft_reference": session.get("draft_reference", ""),
                "prepared_fields": session.get("prepared_fields_json", {}),
                "validation": validation,
                "receipt": receipt,
            },
        )
        updated_application = self.store.update_application(
            application["application_id"],
            {
                "status": "submitted",
                "package": application.get("package_json", {}),
                "review": application.get("review_json", {}),
                "exports": application.get("export_json", {}),
            },
        )
        self.store.record_audit_event(approved_by, "funding.submission.submitted", {**receipt, "governance_ticket_id": governance_ticket_id}, company_id=application["company_id"], application_id=application["application_id"])
        return {"session": updated_session, "application": updated_application, "receipt": receipt, "governance_ticket_id": governance_ticket_id}

    def get_submission_receipt(self, session_id: str) -> dict[str, Any]:
        session = self._require_submission_session(session_id)
        return {"session_id": session_id, "receipt": session.get("receipt_json", {})}

    def list_submission_sessions(self, application_id: str | None = None) -> dict[str, Any]:
        return {"sessions": self.store.list_submission_sessions(application_id)}

    def list_submission_approvals(self, application_id: str | None = None, session_id: str | None = None) -> dict[str, Any]:
        return {"approvals": self.store.list_approval_events(application_id=application_id, session_id=session_id)}

    def list_crm_applications(self, company_id: str = "default") -> dict[str, Any]:
        applications = self.store.list_applications(company_id)
        return {"applications": applications}

    def list_deadlines(self, company_id: str = "default") -> dict[str, Any]:
        calls = self.store.list_calls()
        applications = self.store.list_applications(company_id)
        deadlines = []
        for call in calls:
            if call.get("closes_at"):
                deadlines.append(
                    {
                        "type": "call_deadline",
                        "call_id": call["call_id"],
                        "label": call["title"],
                        "due_at": call["closes_at"],
                        "status": "open",
                    }
                )
        for app in applications:
            session_rows = [item for item in self.store.list_approval_events(application_id=app["application_id"]) if item.get("status") == "pending"]
            if session_rows:
                deadlines.append(
                    {
                        "type": "approval_needed",
                        "application_id": app["application_id"],
                        "label": f"Approval required for {app['application_id']}",
                        "due_at": app["updated_at"],
                        "status": "pending",
                    }
                )
        deadlines.sort(key=lambda item: item.get("due_at") or 9e18)
        return {"deadlines": deadlines}

    def list_alerts(self, company_id: str = "default") -> dict[str, Any]:
        alerts = []
        readiness = self.get_company_readiness(company_id)
        if readiness["missing_fields"]:
            alerts.append(
                {
                    "kind": "missing_company_profile_fields",
                    "severity": "warning",
                    "message": f"Missing company profile fields: {', '.join(readiness['missing_fields'][:5])}",
                    "due_at": None,
                    "application_id": "",
                    "is_resolved": False,
                }
            )
        if readiness["missing_documents"]:
            alerts.append(
                {
                    "kind": "missing_company_documents",
                    "severity": "warning",
                    "message": f"Missing company documents: {', '.join(readiness['missing_documents'])}",
                    "due_at": None,
                    "application_id": "",
                    "is_resolved": False,
                }
            )
        for app in self.store.list_applications(company_id):
            compliance = app.get("package_json", {}).get("compliance", {})
            if compliance.get("missing_documents"):
                alerts.append(
                    {
                        "kind": "application_missing_documents",
                        "severity": "critical",
                        "message": f"Application {app['application_id']} is missing documents: {', '.join(compliance['missing_documents'])}",
                        "due_at": app.get("updated_at"),
                        "application_id": app["application_id"],
                        "is_resolved": False,
                    }
                )
            if app.get("status") == "reviewed":
                alerts.append(
                    {
                        "kind": "ready_for_submission",
                        "severity": "info",
                        "message": f"Application {app['application_id']} is reviewed and ready for submission preparation.",
                        "due_at": app.get("updated_at"),
                        "application_id": app["application_id"],
                        "is_resolved": False,
                    }
                )
        # Listing alerts is called during dashboard route crawls and must remain
        # read-only. Persisting derived alerts here used to turn a GET request
        # into a write path and could lock the shared SQLite database under
        # parallel operator-surface probes.
        return {"alerts": alerts}

    def get_executive_report(self, company_id: str = "default") -> dict[str, Any]:
        profile = self.get_company_profile(company_id)
        readiness = self.get_company_readiness(company_id)
        projects = self.store.list_projects(company_id)
        applications = self.store.list_applications(company_id)
        best_match = None
        if projects:
            latest_project_id = projects[0]["project_id"]
            matches = self.store.list_matches(latest_project_id)
            if matches:
                best_match = matches[0]
        return {
            "company_id": company_id,
            "company_name": profile.get("legal_name", ""),
            "readiness_score": readiness["readiness_score"],
            "open_projects": len(projects),
            "applications": len(applications),
            "best_match": best_match,
            "top_risks": readiness["recommended_next_steps"][:3] + (best_match.get("risks", [])[:3] if best_match else []),
            "generated_at": _now(),
        }
